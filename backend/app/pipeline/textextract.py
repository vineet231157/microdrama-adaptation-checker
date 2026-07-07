"""Extract plain text from an uploaded script (PDF / DOCX / TXT).

Used for the Hindi OG script (Step 5) and for the standalone Formatter (Model 4)
when the user uploads a DOCX/TXT instead of a PDF.
"""
from __future__ import annotations

from pathlib import Path


def extract_text(path: str | Path) -> str:
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _from_pdf(path)
    if ext == ".docx":
        return _from_docx(path)
    # .txt and anything else → read as utf-8 text
    return path.read_text(encoding="utf-8", errors="ignore")


def _from_pdf(path: Path) -> str:
    import pdfplumber
    with pdfplumber.open(str(path)) as pdf:
        return "\n".join((p.extract_text() or "") for p in pdf.pages)


def _from_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    # include table cell text too (scripts sometimes use tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)
