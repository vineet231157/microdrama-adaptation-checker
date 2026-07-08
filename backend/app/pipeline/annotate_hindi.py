"""Annotated Hindi Script (.docx) — the Hindi OG screenplay reproduced verbatim
with two boxes at the top of every episode:

  • ADAPTATION CHANGES  (grey box)  — what the Hindi did differently, incl. anything added
  • MISSING INFORMATION (red box)   — genuine story info a viewer never receives

Python-docx port of the original build_script.js. Boxes are shaded + bordered
via oxml (python-docx has no direct paragraph shading/border API).
"""
from __future__ import annotations

import re
from pathlib import Path

NAVY, GREY, DGREY, RED = "1F3864", "666666", "555555", "B22222"
GREY_FILL, RED_FILL = "F0F0F0", "FDECEA"

_SLUG = re.compile(r"^(Scene\s+\d+[^/]*/\s*)?(INT\.|EXT\.|INT\./EXT\.|MONTAGE)", re.I)
_SLUG2 = re.compile(r"^\d+(-\d+)?\s*/\s*(INT|EXT|MONTAGE)", re.I)
_EP = re.compile(r"^EPISODE\s+(\d+)\b", re.I)
_EP2 = re.compile(r"^EP\s*0?(\d+)\b", re.I)


def _is_slug(l: str) -> bool:
    return bool(_SLUG.match(l) or _SLUG2.match(l))


def _is_cue(l: str) -> bool:
    if len(l) > 46:
        return False
    m = re.match(r"^([A-Z0-9][A-Z0-9 .,'’&/\-]*?)(\s*\([^)]*\))?$", l)
    if not m:
        return False
    name = m.group(1).strip()
    if len(name) < 2 or re.match(r"^(SETTING|SCENE PROFILE)", name):
        return False
    return bool(re.search(r"[A-Z]{2,}", name)) and name == name.upper()


def _starts_lower(l: str) -> bool:
    return bool(re.match(r"^[a-z'\"“(]", l))


# ── oxml shading + border helpers ────────────────────────────────────────────
def _shade(paragraph, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _borders(paragraph, color, top=False, bottom=False, left=True, right=True):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge, on in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        if on:
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "6")
            e.set(qn("w:space"), "2")
            e.set(qn("w:color"), color)
            pBdr.append(e)
    pPr.append(pBdr)


def _run(p, text, *, size=10, bold=False, italic=False, color=None):
    from docx.shared import Pt, RGBColor
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


def _box(doc, label, label_color, fill, items):
    from docx.shared import Pt
    # header
    h = doc.add_paragraph()
    h.paragraph_format.left_indent = Pt(6)
    h.paragraph_format.right_indent = Pt(6)
    h.paragraph_format.space_after = Pt(1)
    _shade(h, fill)
    _borders(h, label_color, top=True, left=True, right=True)
    _run(h, label, size=9, bold=True, color=label_color)
    # bullets
    for i, it in enumerate(items):
        last = i == len(items) - 1
        b = doc.add_paragraph()
        b.paragraph_format.left_indent = Pt(15)
        b.paragraph_format.right_indent = Pt(6)
        b.paragraph_format.space_after = Pt(0 if last else 2)
        _shade(b, fill)
        _borders(b, label_color, bottom=last, left=True, right=True)
        _run(b, "•  " + it, size=8.5, color="222222")


def build(hindi_text: str, out_docx: str, *, title: str, hindi_ann: dict[int, dict]) -> str:
    """hindi_ann: {episode_number: {'added':[], 'gaps':[], 'changes':[]}}."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    # title + legend
    t = doc.add_paragraph(); _run(t, title, size=17, bold=True, color=NAVY)
    s = doc.add_paragraph(); _run(s, "Annotated Hindi Script — adaptation changes & missing "
                                    "information flagged per episode (dialogue verbatim).",
                                  size=11, color=GREY)
    leg = doc.add_paragraph(); _run(leg, "HOW TO READ THIS DOCUMENT", size=10, bold=True, color=NAVY)
    l1 = doc.add_paragraph()
    _run(l1, "•  "); _run(l1, "ADAPTATION CHANGES", bold=True, color=DGREY)
    _run(l1, " (grey box) — everything the Hindi did differently, including anything it added.")
    l2 = doc.add_paragraph()
    _run(l2, "•  "); _run(l2, "MISSING INFORMATION", bold=True, color=RED)
    _run(l2, " (red box) — genuine story info a viewer never receives; says so if nothing is missing.")

    lines = [ln.rstrip() for ln in hindi_text.split("\n")]
    started = False
    buf: list[str] = []

    def flush():
        nonlocal buf
        if buf:
            para = doc.add_paragraph()
            _run(para, " ".join(buf).strip(), size=10)
            buf = []

    def start_episode(n: int):
        nonlocal started
        flush()
        if started:
            doc.add_page_break()
        started = True
        hp = doc.add_heading(f"Episode {n} (Hindi)", level=1)
        for r in hp.runs:
            r.font.color.rgb = RGBColor.from_string(NAVY)
        a = hindi_ann.get(n, {})
        change_list = ["Added — " + x for x in (a.get("added") or [])] + (a.get("changes") or [])
        _box(doc, "ADAPTATION CHANGES", DGREY, GREY_FILL,
             change_list or ["No notable changes."])
        doc.add_paragraph()
        _box(doc, "MISSING INFORMATION", RED, RED_FILL,
             (a.get("gaps") or ["None — no information missing."]))
        doc.add_paragraph()

    for raw in lines:
        l = raw.strip()
        m = _EP.match(l) or _EP2.match(l)
        if m:
            start_episode(int(m.group(1)))
            continue
        if not started or l == "":
            continue
        if re.match(r"^SCENE PROFILE$", l, re.I):
            flush(); continue
        if _is_slug(l):
            flush()
            p = doc.add_paragraph(); _run(p, l, size=10.5, bold=True, color=NAVY)
            continue
        if re.match(r"^SETTING:", l, re.I) or l.startswith("•"):
            flush()
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(8)
            _run(p, l, size=9, italic=True, color=GREY)
            continue
        if _is_cue(l):
            flush()
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(120)
            _run(p, l, size=10, bold=True)
            continue
        if buf and _starts_lower(l):
            buf.append(l)
        else:
            flush(); buf = [l]
    flush()

    # If the Hindi text had NO episode markers, still surface annotations up top.
    if not started and hindi_ann:
        for n in sorted(hindi_ann):
            start_episode(n)
        doc.add_paragraph()
        _run(doc.add_paragraph(), hindi_text[:20000], size=10)

    doc.save(out_docx)
    return out_docx
